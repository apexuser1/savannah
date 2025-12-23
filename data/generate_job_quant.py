#!/usr/bin/env python3
"""
Generate Quantitative Developer Job Description DOCX
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_job_description():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Quantitative Developer / Financial Software Engineer', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(26, 26, 26)
    
    # Company name
    company = doc.add_paragraph('QuantEdge Capital')
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_run = company.runs[0]
    company_run.font.size = Pt(14)
    company_run.font.bold = True
    company_run.font.color.rgb = RGBColor(37, 99, 235)
    
    doc.add_paragraph()
    
    # About the Company
    doc.add_heading('About QuantEdge Capital', level=2)
    about = doc.add_paragraph(
        "QuantEdge Capital is a premier quantitative trading firm and technology-driven hedge fund "
        "managing over $8 billion in assets. Founded in 2010, we leverage cutting-edge mathematics, "
        "statistics, and computer science to develop sophisticated trading strategies across equities, "
        "futures, options, and cryptocurrencies. Our team of 200+ engineers, researchers, and traders "
        "operates from offices in New York, Chicago, London, and Singapore. We pride ourselves on our "
        "collaborative culture, continuous learning environment, and commitment to technological excellence."
    )
    about.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Position Overview
    doc.add_heading('Position Overview', level=2)
    overview = doc.add_paragraph(
        "We are seeking an exceptional Quantitative Developer to join our Trading Systems team. In this "
        "role, you will design, implement, and optimize low-latency trading systems, real-time risk "
        "management platforms, and market data infrastructure. You will work closely with quantitative "
        "researchers and traders to transform mathematical models into production-grade trading systems "
        "that process millions of market events per second with microsecond-level latency requirements. "
        "This is a high-impact role where your code directly influences trading decisions and firm profitability."
    )
    overview.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Key Responsibilities
    doc.add_heading('Key Responsibilities', level=2)
    responsibilities = [
        "Design and implement high-performance, low-latency trading systems in Python and C++",
        
        "Build real-time market data processing infrastructure handling tick data from multiple exchanges",
        
        "Develop order management systems (OMS) and execution management systems (EMS) with sub-millisecond "
        "latency requirements",
        
        "Implement trading algorithms and strategies based on quantitative research models",
        
        "Build real-time risk management and position monitoring systems with automated alerts",
        
        "Optimize database queries and data pipelines for time-series financial data (PostgreSQL, TimescaleDB, kdb+)",
        
        "Integrate with market data APIs (Bloomberg, Reuters, IEX, Polygon.io) and exchange FIX protocols",
        
        "Develop backtesting frameworks and simulation environments for strategy validation",
        
        "Implement monitoring and alerting for production trading systems (24/7 availability critical)",
        
        "Collaborate with quantitative researchers to translate mathematical models into efficient code",
        
        "Perform code reviews and contribute to system architecture decisions"
    ]
    
    for item in responsibilities:
        p = doc.add_paragraph(item, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Required Skills and Qualifications
    doc.add_heading('Required Skills and Qualifications', level=2)
    
    required = [
        ("Programming Languages:", "Expert-level Python and C++ (or Java). Strong understanding of "
         "memory management, concurrency, and performance optimization"),
        
        ("Financial Markets:", "4+ years of experience in financial software development with strong "
         "understanding of equity, futures, or options markets. Knowledge of market microstructure, "
         "order types, and trading mechanics"),
        
        ("Trading Systems:", "Experience building or maintaining trading systems, order management "
         "systems, or market data infrastructure"),
        
        ("Databases:", "Advanced SQL skills with PostgreSQL or similar RDBMS. Experience with time-series "
         "databases for financial data (TimescaleDB, InfluxDB, kdb+)"),
        
        ("Real-time Processing:", "Experience with real-time data processing, event-driven architectures, "
         "and stream processing frameworks"),
        
        ("Low-latency Systems:", "Understanding of low-latency programming techniques, network optimization, "
         "and performance profiling"),
        
        ("Data Structures & Algorithms:", "Strong foundation in computer science fundamentals, with ability "
         "to optimize code for time and space complexity"),
        
        ("Linux/Unix:", "Proficiency with Linux development environment, shell scripting, and system tools"),
        
        ("Experience:", "4-6 years in quantitative development, algorithmic trading, or financial software "
         "engineering"),
        
        ("Education:", "Bachelor's or Master's degree in Computer Science, Mathematics, Financial Engineering, "
         "Physics, or related quantitative field")
    ]
    
    for title, content in required:
        p = doc.add_paragraph(style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        title_run = p.add_run(title)
        title_run.bold = True
        p.add_run(f" {content}")
    
    # Nice to Have
    doc.add_heading('Nice to Have', level=2)
    
    nice_to_have = [
        "Machine learning experience for alpha research, prediction models, or pattern recognition",
        
        "Options pricing models and derivatives knowledge (Black-Scholes, Greeks, volatility surfaces)",
        
        "Risk management systems and portfolio optimization experience",
        
        "FIX protocol expertise and exchange connectivity integration",
        
        "Experience with market data APIs and vendor feeds (Bloomberg API, Reuters, IEX Cloud)",
        
        "Distributed systems and microservices architecture",
        
        "Message queues and pub/sub systems (Kafka, RabbitMQ, ZeroMQ)",
        
        "Performance optimization and profiling tools (gprof, Valgrind, perf)",
        
        "Quantitative research background or advanced degree (MS/PhD) in quantitative field",
        
        "Knowledge of regulatory requirements (SEC, FINRA, MiFID II)"
    ]
    
    for item in nice_to_have:
        p = doc.add_paragraph(item, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Technical Environment
    doc.add_heading('Technical Environment', level=2)
    tech_env = doc.add_paragraph(
        "Python 3.11+, C++17/20, PostgreSQL, TimescaleDB, Redis, Kafka, Docker, Kubernetes, "
        "Git, Linux (RHEL/Ubuntu), Bloomberg API, FIX protocol, REST/WebSocket APIs, "
        "Pandas, NumPy, pytest, gtest"
    )
    tech_env.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Compensation and Benefits
    doc.add_heading('Compensation and Benefits', level=2)
    
    benefits = [
        "Base salary: $180,000 - $250,000 (depending on experience and qualifications)",
        
        "Performance-based bonus: 50-100% of base salary (tied to individual and firm performance)",
        
        "Sign-on bonus for exceptional candidates",
        
        "Comprehensive health insurance (medical, dental, vision) - 100% premium covered",
        
        "401(k) with 8% company contribution (no vesting period)",
        
        "Relocation assistance for candidates moving to New York or Chicago",
        
        "Daily catered lunch and dinner for employees working late",
        
        "Professional development: Conference attendance, online courses, certification reimbursement",
        
        "State-of-the-art workstations (multi-monitor setup, choice of Linux/Mac)",
        
        "Collaborative workspace with game room, gym, and meditation areas",
        
        "20 days PTO + 10 paid holidays + holiday closure between Christmas and New Year",
        
        "Quarterly team outings and annual company retreat"
    ]
    
    for item in benefits:
        p = doc.add_paragraph(item, style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Work Environment
    doc.add_heading('Work Environment & Culture', level=2)
    culture = doc.add_paragraph(
        "QuantEdge values intellectual curiosity, collaboration, and continuous improvement. Our engineering "
        "culture emphasizes code quality, peer learning, and innovation. We hold weekly tech talks, quarterly "
        "hackathons, and maintain an open environment where ideas from all levels are encouraged. While we work "
        "on challenging problems with high-stakes outcomes, we maintain work-life balance and support professional "
        "growth through mentorship programs and internal mobility opportunities."
    )
    culture.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Location
    doc.add_heading('Location', level=2)
    location = doc.add_paragraph(
        "New York, NY or Chicago, IL (On-site required - trading floor environment)"
    )
    location.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    loc_note = doc.add_paragraph(
        "Note: Due to the nature of trading operations and real-time collaboration requirements, "
        "this position requires on-site presence. Remote work is not available for this role."
    )
    loc_note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    loc_note_run = loc_note.runs[0]
    loc_note_run.italic = True
    loc_note_run.font.size = Pt(10)
    
    # How to Apply
    doc.add_heading('How to Apply', level=2)
    apply = doc.add_paragraph(
        "Please submit your application including:\n"
        "1. Resume/CV\n"
        "2. GitHub profile or code samples demonstrating your technical skills\n"
        "3. Brief cover letter explaining your interest in quantitative finance and relevant experience\n\n"
        "Send to: careers@quantedgecapital.com with subject line 'Quantitative Developer Application - [Your Name]'"
    )
    apply.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    # Equal Opportunity
    equal_opp = doc.add_paragraph(
        "QuantEdge Capital is an equal opportunity employer. We are committed to building a diverse and "
        "inclusive workplace and do not discriminate on the basis of race, color, religion, gender, sexual "
        "orientation, national origin, age, disability, veteran status, or any other protected characteristic. "
        "All employment decisions are based on qualifications, merit, and business needs."
    )
    equal_opp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    equal_opp_run = equal_opp.runs[0]
    equal_opp_run.italic = True
    equal_opp_run.font.size = Pt(10)
    
    # Save document
    output_path = "/home/ubuntu/resume_job_matcher/test_data/job_quantitative_developer.docx"
    doc.save(output_path)
    print(f"✅ Created: {output_path}")

if __name__ == "__main__":
    create_job_description()
