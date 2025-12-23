#!/usr/bin/env python3
"""
Generate Fintech Candidate Resume DOCX
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_bullet(doc, text, indent=0.25):
    """Helper function to add a bullet point"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(indent)
    return p

def create_resume():
    doc = Document()
    
    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Name
    name = doc.add_heading('EMILY RODRIGUEZ', level=1)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name.runs[0]
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = RGBColor(26, 26, 26)
    
    # Contact Info
    contact = doc.add_paragraph(
        'New York, NY | (212) 555-0189 | emily.rodriguez@email.com | '
        'linkedin.com/in/emilyrodriguez | github.com/erodriguez-quant'
    )
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact.runs[0]
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = RGBColor(75, 85, 99)
    
    doc.add_paragraph()
    
    # Professional Summary
    doc.add_heading('PROFESSIONAL SUMMARY', level=2)
    summary = doc.add_paragraph(
        "Highly skilled Quantitative Developer with 5+ years of experience building high-performance trading "
        "systems, real-time market data infrastructure, and quantitative analytics platforms. Expert in Python "
        "and C++ with deep understanding of financial markets, low-latency systems, and algorithmic trading. "
        "Proven track record of optimizing trading systems to achieve sub-millisecond execution latency and "
        "processing 100M+ market data events daily. Strong foundation in computer science, mathematics, and "
        "financial engineering. Passionate about leveraging technology to solve complex quantitative problems."
    )
    summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Technical Skills
    doc.add_heading('TECHNICAL SKILLS', level=2)
    
    skills_table_data = [
        ("Programming Languages:", "Python (expert), C++ (advanced), SQL, Bash scripting"),
        ("Financial Systems:", "Trading systems, Order management systems (OMS), Real-time market data processing, "
         "Risk management platforms"),
        ("Databases:", "PostgreSQL, TimescaleDB (time-series), Redis, MongoDB, kdb+ (basic)"),
        ("Financial Markets:", "Equities, Futures, Options, Market microstructure, Order types, Trading protocols"),
        ("Data Processing:", "Real-time stream processing, Event-driven architectures, Message queues (Kafka, RabbitMQ)"),
        ("APIs & Protocols:", "REST APIs, WebSocket APIs, FIX protocol, Bloomberg API, Interactive Brokers API"),
        ("Development Tools:", "Git, Docker, Linux/Unix, pytest, gtest, Valgrind, gprof"),
        ("Libraries & Frameworks:", "Pandas, NumPy, asyncio, FastAPI, Flask, Boost (C++)"),
        ("Market Data:", "Bloomberg Terminal, Reuters, IEX Cloud, Polygon.io, Alpha Vantage")
    ]
    
    # Create a simple table-like structure with tabs
    for label, content in skills_table_data:
        p = doc.add_paragraph()
        label_run = p.add_run(label)
        label_run.bold = True
        p.add_run(f" {content}")
        p.paragraph_format.left_indent = Inches(0)
    
    # Professional Experience
    doc.add_heading('PROFESSIONAL EXPERIENCE', level=2)
    
    # Job 1
    job1_title = doc.add_paragraph()
    title_run = job1_title.add_run('Senior Quantitative Developer')
    title_run.bold = True
    title_run.font.size = Pt(12)
    date_run = job1_title.add_run(' ' * 50 + 'March 2021 - Present')
    date_run.bold = True
    date_run.font.size = Pt(11)
    
    company1 = doc.add_paragraph('AlphaQuant Trading, New York, NY')
    company1_run = company1.runs[0]
    company1_run.italic = True
    
    job1_bullets = [
        "Designed and implemented low-latency trading execution system in C++ achieving average execution "
        "latency of 450 microseconds, reducing slippage by 30% and improving P&L by $2M annually",
        
        "Built real-time market data processing pipeline in Python handling 100M+ tick updates daily from "
        "multiple exchanges (NYSE, NASDAQ, CME) with fault-tolerant architecture",
        
        "Developed order management system (OMS) supporting multiple order types (limit, market, stop, "
        "iceberg) with pre-trade risk checks and post-trade analytics",
        
        "Implemented real-time risk management platform monitoring position limits, P&L, Greeks exposure, "
        "and margin requirements with automated alert system",
        
        "Optimized database queries and data models in PostgreSQL/TimescaleDB, reducing query time for "
        "historical analysis from minutes to seconds",
        
        "Integrated FIX protocol connectivity to multiple brokers and exchanges for order routing and "
        "execution reporting",
        
        "Built backtesting framework allowing researchers to validate strategies on 10+ years of historical "
        "tick data with realistic execution simulation",
        
        "Collaborated with quantitative researchers to productionize alpha signals and trading strategies, "
        "reducing time-to-production from weeks to days",
        
        "Mentored 2 junior developers on financial systems architecture, C++ optimization, and trading "
        "system best practices"
    ]
    
    for bullet in job1_bullets:
        add_bullet(doc, bullet)
    
    doc.add_paragraph()
    
    # Job 2
    job2_title = doc.add_paragraph()
    title_run = job2_title.add_run('Quantitative Software Engineer')
    title_run.bold = True
    title_run.font.size = Pt(12)
    date_run = job2_title.add_run(' ' * 45 + 'June 2019 - February 2021')
    date_run.bold = True
    date_run.font.size = Pt(11)
    
    company2 = doc.add_paragraph('QuantumEdge Capital, Chicago, IL')
    company2_run = company2.runs[0]
    company2_run.italic = True
    
    job2_bullets = [
        "Developed high-performance options pricing engine in C++ implementing Black-Scholes, "
        "binomial models, and implied volatility calculation with GPU acceleration",
        
        "Built Python-based analytics platform for portfolio risk analysis, calculating VaR, Greeks, "
        "and scenario analysis for $500M+ portfolio",
        
        "Implemented real-time market data feeds integration with Bloomberg API and Reuters for "
        "equity and derivatives pricing",
        
        "Created automated trading strategies execution system supporting pairs trading, statistical "
        "arbitrage, and mean reversion strategies",
        
        "Designed time-series database schema in TimescaleDB for efficient storage and retrieval of "
        "tick data, reducing storage costs by 40%",
        
        "Developed monitoring dashboards and alerting system for trading system health, execution "
        "quality, and strategy performance",
        
        "Optimized critical path code achieving 60% performance improvement through profiling (Valgrind, "
        "gprof) and algorithmic optimization"
    ]
    
    for bullet in job2_bullets:
        add_bullet(doc, bullet)
    
    doc.add_paragraph()
    
    # Job 3
    job3_title = doc.add_paragraph()
    title_run = job3_title.add_run('Junior Quantitative Developer')
    title_run.bold = True
    title_run.font.size = Pt(12)
    date_run = job3_title.add_run(' ' * 43 + 'July 2018 - May 2019')
    date_run.bold = True
    date_run.font.size = Pt(11)
    
    company3 = doc.add_paragraph('FinTech Innovations Inc., New York, NY')
    company3_run = company3.runs[0]
    company3_run.italic = True
    
    job3_bullets = [
        "Developed Python tools for market data analysis, including volatility surface construction and "
        "correlation analysis",
        
        "Built automated data validation and quality checks for incoming market data feeds, reducing "
        "data errors by 85%",
        
        "Implemented REST APIs for trading system integration using Flask and FastAPI",
        
        "Created PostgreSQL database schemas and ETL pipelines for historical market data storage",
        
        "Participated in on-call rotation supporting production trading systems with 24/7 availability"
    ]
    
    for bullet in job3_bullets:
        add_bullet(doc, bullet)
    
    # Education
    doc.add_heading('EDUCATION', level=2)
    
    edu1_title = doc.add_paragraph()
    degree_run = edu1_title.add_run('Master of Science in Financial Engineering')
    degree_run.bold = True
    degree_run.font.size = Pt(11)
    date_run = edu1_title.add_run(' ' * 30 + '2016 - 2018')
    date_run.bold = True
    
    school1 = doc.add_paragraph('Columbia University, New York, NY')
    school1_run = school1.runs[0]
    school1_run.italic = True
    
    details1 = doc.add_paragraph(
        'GPA: 3.8/4.0 | Relevant Coursework: Derivatives Pricing, Algorithmic Trading, Stochastic Calculus, '
        'Machine Learning for Finance, High-Frequency Trading, Risk Management'
    )
    
    doc.add_paragraph()
    
    edu2_title = doc.add_paragraph()
    degree_run = edu2_title.add_run('Bachelor of Science in Computer Science')
    degree_run.bold = True
    degree_run.font.size = Pt(11)
    date_run = edu2_title.add_run(' ' * 35 + '2012 - 2016')
    date_run.bold = True
    
    school2 = doc.add_paragraph('University of Illinois at Urbana-Champaign')
    school2_run = school2.runs[0]
    school2_run.italic = True
    
    details2 = doc.add_paragraph(
        'GPA: 3.7/4.0 | Minor in Mathematics | Relevant Coursework: Data Structures, Algorithms, '
        'Operating Systems, Database Systems, Distributed Systems'
    )
    
    # Key Projects
    doc.add_heading('KEY PROJECTS & ACHIEVEMENTS', level=2)
    
    projects = [
        ("High-Frequency Market Making System:", "Led development of market making algorithm in C++ that "
         "quotes bid-ask spreads for equities, generating $500K profit in first quarter with 65% fill rate"),
        
        ("Low-Latency Order Router:", "Built multi-exchange smart order router optimizing execution across "
         "venues based on liquidity, latency, and fees, reducing execution costs by 15 bps"),
        
        ("Machine Learning Alpha Research:", "Developed Python framework for ML-based alpha signal research "
         "using gradient boosting and neural networks, identifying 3 profitable signals in production"),
        
        ("Real-Time Risk Dashboard:", "Created web-based risk monitoring dashboard with WebSocket updates "
         "showing live P&L, positions, and Greeks exposure used by entire trading desk")
    ]
    
    for title, description in projects:
        p = doc.add_paragraph()
        title_run = p.add_run(title)
        title_run.bold = True
        p.add_run(f" {description}")
        p.paragraph_format.left_indent = Inches(0.25)
    
    # Publications & Certifications
    doc.add_heading('CERTIFICATIONS & PROFESSIONAL DEVELOPMENT', level=2)
    
    certs = [
        "CFA Level II Candidate (exam scheduled June 2024)",
        "Completed 'Advanced C++ Programming' course - LinkedIn Learning (2022)",
        "Attended QuantCon 2023 - Quantitative Trading Conference",
        "Member of International Association for Quantitative Finance (IAQF)"
    ]
    
    for cert in certs:
        add_bullet(doc, cert)
    
    # Save document
    output_path = "/home/ubuntu/resume_job_matcher/test_data/resume_fintech_candidate.docx"
    doc.save(output_path)
    print(f"✅ Created: {output_path}")

if __name__ == "__main__":
    create_resume()
